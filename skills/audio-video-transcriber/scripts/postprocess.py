#!/usr/bin/env python3
"""Create agent-ready review deliverables for a transcript."""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import html
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path


DEFAULT_BASE_DIR = Path("~/AudioVideoTranscriber").expanduser()
TIMESTAMP_RE = re.compile(
    r"(?P<start>\d{1,2}:\d{2}:\d{2}[,.]\d{1,3})\s*-->\s*(?P<end>\d{1,2}:\d{2}:\d{2}[,.]\d{1,3})"
)


@dataclass(frozen=True)
class TranscriptSegment:
    text: str
    start: str = ""
    end: str = ""


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    rows: tuple[tuple[str, ...], ...] = ()
    items: tuple[str, ...] = ()


def expand_path(value: str | Path) -> Path:
    return Path(value).expanduser()


def default_output_dir() -> Path:
    base_dir = expand_path(os.environ.get("AVT_BASE_DIR", str(DEFAULT_BASE_DIR)))
    return base_dir / "output"


def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def clean_transcript(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.fullmatch(r"\d+", line):
            continue
        if "-->" in line:
            continue
        if line.upper() == "WEBVTT":
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def format_seconds(value: object) -> str:
    try:
        seconds = float(value)
    except (TypeError, ValueError):
        return ""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"


def format_milliseconds(value: str) -> str:
    try:
        seconds = int(value) / 1000
    except (TypeError, ValueError):
        return ""
    return format_seconds(seconds)


def parse_timed_text(raw_text: str) -> list[TranscriptSegment]:
    segments: list[TranscriptSegment] = []
    current_start = ""
    current_end = ""
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer, current_start, current_end
        text = " ".join(line.strip() for line in buffer if line.strip()).strip()
        if text:
            segments.append(TranscriptSegment(text=text, start=current_start, end=current_end))
        buffer = []

    for raw_line in raw_text.splitlines():
        line = raw_line.strip()
        if not line or line.upper() == "WEBVTT" or re.fullmatch(r"\d+", line):
            if not line:
                flush()
            continue
        timestamp = TIMESTAMP_RE.search(line)
        if timestamp:
            flush()
            current_start = timestamp.group("start").replace(",", ".")
            current_end = timestamp.group("end").replace(",", ".")
            continue
        buffer.append(line)
    flush()
    return segments


def read_transcript_segments(path: Path) -> list[TranscriptSegment]:
    raw_text = read_text(path)
    suffix = path.suffix.lower()
    if suffix == ".json":
        try:
            data = json.loads(raw_text)
        except json.JSONDecodeError:
            return [TranscriptSegment(text=clean_transcript(raw_text))]
        if isinstance(data, dict):
            segments = data.get("segments")
            if isinstance(segments, list):
                parsed = [
                    TranscriptSegment(
                        text=str(segment.get("text", "")).strip(),
                        start=format_seconds(segment.get("start")),
                        end=format_seconds(segment.get("end")),
                    )
                    for segment in segments
                    if isinstance(segment, dict) and str(segment.get("text", "")).strip()
                ]
                if parsed:
                    return parsed
            if isinstance(data.get("text"), str):
                return paragraphs_to_segments(data["text"])
        return [TranscriptSegment(text=clean_transcript(raw_text))]
    if suffix == ".tsv":
        rows = csv.DictReader(raw_text.splitlines(), delimiter="\t")
        if rows.fieldnames and "text" in rows.fieldnames:
            parsed = [
                TranscriptSegment(
                    text=row.get("text", "").strip(),
                    start=format_milliseconds(row.get("start", "")),
                    end=format_milliseconds(row.get("end", "")),
                )
                for row in rows
                if row.get("text", "").strip()
            ]
            if parsed:
                return parsed
        text_rows = []
        for line in raw_text.splitlines():
            parts = line.split("\t")
            if parts and parts[-1].strip().lower() != "text":
                text_rows.append(parts[-1].strip())
        return paragraphs_to_segments("\n".join(line for line in text_rows if line))
    if suffix in {".srt", ".vtt"}:
        parsed = parse_timed_text(raw_text)
        if parsed:
            return parsed
    return paragraphs_to_segments(clean_transcript(raw_text))


def paragraphs_to_segments(text: str) -> list[TranscriptSegment]:
    parts = [part.strip() for part in re.split(r"\n{2,}|\n", text) if part.strip()]
    return [TranscriptSegment(text=part) for part in parts] or [TranscriptSegment(text="")]


def read_transcript(path: Path) -> str:
    return "\n".join(segment.text for segment in read_transcript_segments(path) if segment.text).strip()


def preview_chunks(text: str, max_chunks: int = 8, chunk_chars: int = 240) -> list[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    chunks = []
    for start in range(0, min(len(normalized), max_chunks * chunk_chars), chunk_chars):
        chunk = normalized[start : start + chunk_chars].strip()
        if chunk:
            chunks.append(chunk)
    return chunks


def word_count_rough(text: str) -> int:
    ascii_words = re.findall(r"[A-Za-z0-9_]+", text)
    cjk_chars = re.findall(r"[\u4e00-\u9fff]", text)
    return len(ascii_words) + len(cjk_chars)


def detect_data_signals(text: str, limit: int = 10) -> list[str]:
    patterns = [
        r"[\d,.]+(?:万|亿|千|百|元|美元|人民币|人|个|次|家|%|％)",
        r"\d{4}年\d{1,2}月\d{0,2}日?",
        r"\d{1,2}月\d{1,2}日",
        r"\d{1,2}:\d{2}",
        r"(?:增长|下降|提升|降低|增加|减少|超过|达到|排名|第)\S{0,12}",
    ]
    signals: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            start = max(0, match.start() - 35)
            end = min(len(text), match.end() + 35)
            snippet = re.sub(r"\s+", " ", text[start:end]).strip()
            if snippet and snippet not in signals:
                signals.append(snippet)
            if len(signals) >= limit:
                return signals
    return signals


def write_if_needed(path: Path, content: str | bytes, overwrite: bool) -> bool:
    if path.exists() and not overwrite:
        return False
    if isinstance(content, bytes):
        path.write_bytes(content)
    else:
        path.write_text(content, encoding="utf-8")
    return True


def output_status(path: Path, written: bool) -> str:
    return "written" if written else "exists"


def print_initial_review_guidance(
    transcript_path: Path,
    output_dir: Path,
    stem: str,
    kinds: list[str],
    include_html: bool,
) -> None:
    print("Review draft guidance:")
    transcript_docx = output_dir / f"{stem}.transcript.docx"
    direct_transcript = transcript_docx if transcript_docx.exists() else transcript_path
    print(f"  Directly viewable transcript: {direct_transcript}")
    for kind in kinds:
        label = title_for_kind(kind)
        print(f"  {label} Markdown: {output_dir / markdown_filename(stem, kind)}")
        if include_html:
            print(f"  {label} HTML: {output_dir / f'{stem}.{kind}.html'}")
    print("  These review files are draft shells until an Agent or user completes the Markdown.")
    print("Next step for final deliverables:")
    print("  Fill the generated Markdown, then sync it to Word/HTML with:")
    print(f"  ./bin/avt review-sync \"{transcript_path}\" --all")


def print_transcript_guidance(transcript_path: Path, output_dir: Path, stem: str) -> None:
    transcript_docx = output_dir / f"{stem}.transcript.docx"
    print("Transcript deliverables:")
    print(f"  原始转写稿: {transcript_path}")
    if transcript_docx.exists():
        print(f"  Word 版全文稿: {transcript_docx}")
    print("我已完成原始转写稿。这版保留了口语表达和原始顺序。")
    print("如果需要，我还可以继续整理成：内容总结版、勘误精修版、刊物发布版、会议纪要版、逐字稿清洁版。")


def print_sync_guidance(results: list[tuple[Path, bool]]) -> None:
    docx_files = [path for path, _ in results if path.suffix.lower() == ".docx"]
    html_files = [path for path, _ in results if path.suffix.lower() == ".html"]
    print("Final deliverables:")
    if docx_files:
        print("  Word files for direct handoff:")
        for path in docx_files:
            print(f"    {path}")
    if html_files:
        print("  HTML files for web/Notion/Feishu copy:")
        for path in html_files:
            print(f"    {path}")
    if not docx_files and not html_files:
        print("  No final DOCX/HTML files were written.")


def build_info_lines(transcript_path: Path, source_file: str, generated_at: str, output_formats: list[str]) -> list[str]:
    return [
        f"原始文件名：{Path(source_file).name if source_file else 'unknown'}",
        f"转写/整理时间：{generated_at}",
        f"Transcript 路径：{transcript_path}",
        f"输出格式列表：{', '.join(output_formats)}",
    ]


def build_summary(transcript_path: Path, source_file: str, transcript: str, generated_at: str) -> str:
    chunks = preview_chunks(transcript)
    chunk_lines = "\n".join(f"{idx}. {chunk}" for idx, chunk in enumerate(chunks, 1))
    if not chunk_lines:
        chunk_lines = "1. Transcript is empty or could not be parsed."

    data_rows = "\n".join(f"| 待 agent 分析 | {signal} | 待 agent 填写 | 待 agent 填写 |" for signal in detect_data_signals(transcript))
    if not data_rows:
        data_rows = "| 待 agent 核对 | 待 agent 填写 | 待 agent 填写 | 未发现明显数据时保留此模板 |"

    return f"""# 内容总结

Source media: `{source_file or "unknown"}`

Transcript: `{transcript_path}`

Generated at: {generated_at}

> Agent task: read the transcript at the path above and replace the draft sections with a polished Chinese summary. If no local LLM API is available, keep this as the handoff template for a future agent pass.

## 核心摘要

待 agent 精修：用 1-3 句话概括这段音频/视频最核心的信息。

## 重点内容

待 agent 精修：

- 提炼 5-10 条最重要内容。
- 合并重复表达。
- 保留关键名词、人名、地名、产品名、数据和结论。

## 关键观点

待 agent 精修：提炼说话人表达出的判断、立场、建议、分歧和结论。

## 待办事项

待 agent 精修：如果转写中出现明确任务、承诺、负责人、时间点或下一步，请列在这里。

| 待办 | 负责人 | 截止时间 | 上下文 |
| --- | --- | --- | --- |
| 待 agent 核对 | 待 agent 填写 | 待 agent 填写 | 待 agent 填写 |

## 时间线

待 agent 精修：按出现顺序列出关键事件、话题转换、时间点、决定和结论。

| 时间/顺序 | 内容 | 备注 |
| --- | --- | --- |
| 待 agent 核对 | 待 agent 填写 | - |

## 数据与信息分析

待 agent 精修：提取金额、百分比、日期、时间、数量、排名、对比关系、增长/下降、用户规模、业务指标和会议任务节点，并说明可能含义。

| 数据/信息 | 原文位置或上下文 | 可能含义 | 备注 |
| --- | --- | --- | --- |
{data_rows}

## 精华版总结

待 agent 精修：

- 重要观点：
- 重要例子：
- 可执行建议：
- 值得引用的表达：

## 风险与疑点

待 agent 精修：记录可能听错、逻辑跳跃、需要回听确认的地方。

## 转写线索预览

{chunk_lines}

## Transcript Stats

- Approximate length: {len(transcript)} characters
- Rough token/word signal: {word_count_rough(transcript)}
"""


def build_corrections(transcript_path: Path, source_file: str, transcript: str, generated_at: str) -> str:
    return f"""# 勘误与精修版

Source media: `{source_file or "unknown"}`

Transcript: `{transcript_path}`

Generated at: {generated_at}

> Agent task: read the transcript at the path above, identify likely ASR mistakes, and produce a corrected transcript. If no local LLM API is available, keep this as the handoff template for a future agent pass.

## 疑似识别错误表

| 原句 | 疑似错词 | 建议修正 | 理由/上下文 | 置信度 |
| --- | --- | --- | --- | --- |
| 待 agent 核对 | 待 agent 填写 | 待 agent 填写 | 结合上下文判断 | - |

## 专有名词统一表

| 原句 | 原转写 | 建议修正 | 类型 | 置信度 |
| --- | --- | --- | --- | --- |
| 待 agent 核对 | 待 agent 填写 | 待 agent 填写 | 人名/地名/机构名/品牌名/术语 | - |

## 语句精修说明

| 原句 | 建议修正 | 类型 | 理由/上下文 | 置信度 |
| --- | --- | --- | --- | --- |
| 待 agent 核对 | 待 agent 填写 | 错词/断句/数字/术语/重复 | 结合上下文判断 | - |

## 常见优先核对项

- 人名、公司名、品牌名、地名、楼盘/建筑名
- 英文缩写、产品名、技术词
- 数字、金额、百分比、年份、时间点
- 同音词、近音词、跨语言音译词
- Whisper 明显断句错误或重复片段

## 精修后正文

待 agent 精修：在这里输出修正错别字、术语、断句和段落结构后的版本。保持原意，不要擅自添加转写中没有的信息。

## 原始转写

```text
{transcript}
```
"""


def build_publish(transcript_path: Path, source_file: str, transcript: str, generated_at: str) -> str:
    chunks = preview_chunks(transcript, max_chunks=10, chunk_chars=300)
    source_material = "\n\n".join(chunks) if chunks else transcript or "Transcript is empty or could not be parsed."
    return f"""# 刊物发布版

Source media: `{source_file or "unknown"}`

Transcript: `{transcript_path}`

Generated at: {generated_at}

> Agent task: read the transcript at the path above and turn it into a publishable Chinese article. Preserve the speaker's meaning, keep core viewpoints, add only light transitions when needed, and do not invent facts.

## 发布稿正文

待 agent 精修：在不改变原意的基础上，整理为自然、连贯、有段落结构的文章。可以适度补充衔接语，但不得编造事实。

## 人工审校提示

- 核对人名、机构名、地名、品牌名和术语。
- 核对数字、日期、金额、百分比和引用。
- 删除不适合发布的口头停顿、重复表达和无意义语气词。
- 保留讲话者核心观点，不把推测写成事实。

## 转写素材

```text
{source_material}
```
"""


def build_meeting_notes(transcript_path: Path, source_file: str, transcript: str, generated_at: str) -> str:
    chunks = preview_chunks(transcript, max_chunks=10, chunk_chars=300)
    chunk_lines = "\n".join(f"{idx}. {chunk}" for idx, chunk in enumerate(chunks, 1))
    if not chunk_lines:
        chunk_lines = "1. Transcript is empty or could not be parsed."
    return f"""# 会议纪要版

Source media: `{source_file or "unknown"}`

Transcript: `{transcript_path}`

Generated at: {generated_at}

> Agent task: read the transcript at the path above and create meeting notes for quick reporting. Extract the topic, core viewpoints, action items, risks, and next steps. Do not invent decisions or owners.

## 会议主题

待 agent 整理：用一句话说明这段会议/讨论的主题。

## 核心观点

- 待 agent 提炼。

## 待办事项

| 待办 | 负责人 | 截止时间 | 上下文 |
| --- | --- | --- | --- |
| 待 agent 核对 | 待 agent 填写 | 待 agent 填写 | 待 agent 填写 |

## 风险点

- 待 agent 提炼需要确认、存在分歧或可能影响后续推进的事项。

## 后续动作

- 待 agent 提炼下一步行动。

## 转写线索预览

{chunk_lines}
"""


def build_html_document(title: str, info_lines: list[str], sections: list[tuple[str, str]]) -> str:
    section_html = []
    for heading, body in sections:
        section_html.append(f"<section><h2>{html.escape(heading)}</h2>{body}</section>")
    info_items = "\n".join(f"<li>{html.escape(line)}</li>" for line in info_lines)
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    body {{
      margin: 0;
      color: #1f2933;
      background: #f6f7f9;
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
      line-height: 1.75;
    }}
    main {{
      max-width: 920px;
      margin: 0 auto;
      padding: 40px 24px;
      background: #ffffff;
      min-height: 100vh;
    }}
    header {{
      border-bottom: 3px solid #2f6f73;
      padding-bottom: 20px;
      margin-bottom: 28px;
    }}
    h1 {{ margin: 0 0 10px; font-size: 30px; }}
    h2 {{ margin-top: 32px; padding-bottom: 6px; border-bottom: 1px solid #d7dee6; font-size: 22px; }}
    .info-card {{
      background: #eef6f6;
      border: 1px solid #c9dddd;
      border-radius: 8px;
      padding: 14px 18px;
    }}
    table {{ width: 100%; border-collapse: collapse; margin: 14px 0; }}
    th, td {{ border: 1px solid #d7dee6; padding: 9px 10px; vertical-align: top; }}
    th {{ background: #f0f4f8; font-weight: 700; }}
    .placeholder {{ color: #64748b; }}
    .transcript p {{ margin: 0 0 12px; }}
    .time {{ color: #4b5563; font-weight: 700; }}
  </style>
</head>
<body>
<main>
  <header>
    <h1>{html.escape(title)}</h1>
    <div class="info-card"><ul>{info_items}</ul></div>
  </header>
  {''.join(section_html)}
</main>
</body>
</html>
"""


def paragraph_html(text: str, css_class: str = "placeholder") -> str:
    return f'<p class="{css_class}">{html.escape(text)}</p>'


def table_html(headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{html.escape(header)}</th>" for header in headers)
    body = "\n".join(
        "<tr>" + "".join(f"<td>{html.escape(cell)}</td>" for cell in row) + "</tr>"
        for row in rows
    )
    return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"


def split_markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_markdown_table_separator(line: str) -> bool:
    cells = split_markdown_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def parse_markdown_blocks(markdown_text: str) -> list[MarkdownBlock]:
    blocks: list[MarkdownBlock] = []
    lines = markdown_text.splitlines()
    index = 0
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        text = "\n".join(line.strip() for line in paragraph if line.strip()).strip()
        if text:
            blocks.append(MarkdownBlock(kind="paragraph", text=text))
        paragraph = []

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(MarkdownBlock(kind="code", text="\n".join(code_lines).strip()))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    kind="heading",
                    level=len(heading_match.group(1)),
                    text=heading_match.group(2).strip(),
                )
            )
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_markdown_table_separator(lines[index + 1]):
            flush_paragraph()
            rows = [tuple(split_markdown_row(stripped))]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(tuple(split_markdown_row(lines[index])))
                index += 1
            blocks.append(MarkdownBlock(kind="table", rows=tuple(rows)))
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            items = []
            while index < len(lines):
                item_match = re.match(r"^[-*]\s+(.+)$", lines[index].strip())
                if not item_match:
                    break
                items.append(item_match.group(1).strip())
                index += 1
            blocks.append(MarkdownBlock(kind="list", items=tuple(items)))
            continue

        paragraph.append(stripped.lstrip("> ").strip())
        index += 1

    flush_paragraph()
    return blocks


def markdown_title(markdown_text: str, fallback: str) -> str:
    for block in parse_markdown_blocks(markdown_text):
        if block.kind == "heading" and block.level == 1 and block.text:
            return block.text
    return fallback


def extract_backtick_value(markdown_text: str, label: str) -> str:
    pattern = rf"^{re.escape(label)}:\s*`([^`]+)`\s*$"
    for line in markdown_text.splitlines():
        match = re.match(pattern, line.strip())
        if match:
            return match.group(1).strip()
    return ""


def build_sync_info_lines(markdown_path: Path, markdown_text: str, generated_at: str, output_formats: list[str]) -> list[str]:
    transcript_path = extract_backtick_value(markdown_text, "Transcript") or "unknown"
    source_media = extract_backtick_value(markdown_text, "Source media") or "unknown"
    return [
        f"来源 Markdown：{markdown_path}",
        f"原始文件名：{Path(source_media).name if source_media != 'unknown' else 'unknown'}",
        f"同步时间：{generated_at}",
        f"Transcript 路径：{transcript_path}",
        f"输出格式列表：{', '.join(output_formats)}",
    ]


def markdown_blocks_to_html(markdown_text: str) -> str:
    rendered = []
    for block in parse_markdown_blocks(markdown_text):
        if block.kind == "heading":
            level = min(max(block.level + 1, 2), 6)
            rendered.append(f"<h{level}>{html.escape(block.text)}</h{level}>")
        elif block.kind == "paragraph":
            paragraphs = [part.strip() for part in block.text.splitlines() if part.strip()]
            rendered.extend(f"<p>{html.escape(part)}</p>" for part in paragraphs)
        elif block.kind == "list":
            items = "".join(f"<li>{html.escape(item)}</li>" for item in block.items)
            rendered.append(f"<ul>{items}</ul>")
        elif block.kind == "table" and block.rows:
            headers = list(block.rows[0])
            rows = [list(row) for row in block.rows[1:]]
            rendered.append(table_html(headers, rows))
        elif block.kind == "code":
            rendered.append(f"<pre><code>{html.escape(block.text)}</code></pre>")
    return "\n".join(rendered)


def build_markdown_sync_html(title: str, info_lines: list[str], markdown_text: str) -> str:
    return build_html_document(title, info_lines, [("正文", markdown_blocks_to_html(markdown_text))])


def build_summary_html(info_lines: list[str], transcript: str) -> str:
    data_rows = [["待 agent 核对", signal, "待 agent 填写", "待 agent 填写"] for signal in detect_data_signals(transcript)]
    if not data_rows:
        data_rows = [["待 agent 核对", "待 agent 填写", "待 agent 填写", "未发现明显数据时保留此模板"]]
    sections = [
        ("核心摘要", paragraph_html("待 agent 精修：用 1-3 句话概括最核心的信息。")),
        ("重点内容", paragraph_html("待 agent 精修：提炼 5-10 条最重要内容，合并重复表达。")),
        ("关键观点", paragraph_html("待 agent 精修：提炼判断、立场、建议、分歧和结论。")),
        ("待办事项", table_html(["待办", "负责人", "截止时间", "上下文"], [["待 agent 核对", "待 agent 填写", "待 agent 填写", "待 agent 填写"]])),
        ("时间线", table_html(["时间/顺序", "内容", "备注"], [["待 agent 核对", "待 agent 填写", "-"]])),
        ("数据与信息分析", table_html(["数据/信息", "原文位置或上下文", "可能含义", "备注"], data_rows)),
        ("精华版总结", paragraph_html("待 agent 精修：输出适合直接阅读或交付的精华摘要。")),
    ]
    return build_html_document("内容总结", info_lines, sections)


def build_corrections_html(info_lines: list[str]) -> str:
    sections = [
        ("疑似识别错误表", table_html(["原句", "疑似错词", "建议修正", "理由/上下文", "置信度"], [["待 agent 核对", "待 agent 填写", "待 agent 填写", "结合上下文判断", "-"]])),
        ("专有名词统一表", table_html(["原句", "原转写", "建议修正", "类型", "置信度"], [["待 agent 核对", "待 agent 填写", "待 agent 填写", "人名/地名/机构名/品牌名/术语", "-"]])),
        ("语句精修说明", table_html(["原句", "建议修正", "类型", "理由/上下文", "置信度"], [["待 agent 核对", "待 agent 填写", "错词/断句/数字/术语/重复", "结合上下文判断", "-"]])),
        ("精修后正文", paragraph_html("待 agent 精修：在这里输出修正错别字、术语、断句和段落结构后的版本。")),
    ]
    return build_html_document("勘误与精修版", info_lines, sections)


def import_docx():
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install it with: python -m pip install python-docx"
        ) from exc
    return Document, Pt, qn


def set_doc_style(document: object, pt: object, qn: object) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_info_table(document: object, info_lines: list[str]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for line in info_lines:
        key, _, value = line.partition("：")
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value


def add_simple_table(document: object, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def save_document(document: object, path: Path, overwrite: bool) -> bool:
    if path.exists() and not overwrite:
        return False
    document.save(path)
    return True


def create_transcript_docx(
    path: Path,
    transcript_path: Path,
    source_file: str,
    generated_at: str,
    output_formats: list[str],
    segments: list[TranscriptSegment],
    overwrite: bool,
) -> bool:
    Document, Pt, qn = import_docx()
    document = Document()
    set_doc_style(document, Pt, qn)
    source_name = Path(source_file).stem if source_file else transcript_path.stem
    document.add_heading(f"{source_name} 转写稿", level=0)
    add_info_table(document, build_info_lines(transcript_path, source_file, generated_at, output_formats))
    document.add_heading("完整转写稿", level=1)
    for segment in segments:
        if segment.start or segment.end:
            document.add_paragraph(f"[{segment.start} - {segment.end}]")
        for paragraph in re.split(r"\n{2,}|\n", segment.text):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
    return save_document(document, path, overwrite)


def create_summary_docx(
    path: Path,
    transcript_path: Path,
    source_file: str,
    generated_at: str,
    output_formats: list[str],
    transcript: str,
    overwrite: bool,
) -> bool:
    Document, Pt, qn = import_docx()
    document = Document()
    set_doc_style(document, Pt, qn)
    document.add_heading("内容总结", level=0)
    add_info_table(document, build_info_lines(transcript_path, source_file, generated_at, output_formats))
    sections = [
        ("核心摘要", "待 agent 精修：用 1-3 句话概括这段音频/视频最核心的信息。"),
        ("重点内容", "待 agent 精修：提炼 5-10 条最重要内容，合并重复表达。"),
        ("关键观点", "待 agent 精修：提炼说话人表达出的判断、立场、建议、分歧和结论。"),
    ]
    for heading, body in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(body)
    document.add_heading("待办事项", level=1)
    add_simple_table(document, ["待办", "负责人", "截止时间", "上下文"], [["待 agent 核对", "待 agent 填写", "待 agent 填写", "待 agent 填写"]])
    document.add_heading("时间线", level=1)
    add_simple_table(document, ["时间/顺序", "内容", "备注"], [["待 agent 核对", "待 agent 填写", "-"]])
    document.add_heading("数据与信息分析", level=1)
    data_rows = [["待 agent 分析", signal, "待 agent 填写", "待 agent 填写"] for signal in detect_data_signals(transcript)]
    if not data_rows:
        data_rows = [["待 agent 核对", "待 agent 填写", "待 agent 填写", "未发现明显数据时保留此模板"]]
    add_simple_table(document, ["数据/信息", "原文位置或上下文", "可能含义", "备注"], data_rows)
    document.add_heading("精华版总结", level=1)
    document.add_paragraph("待 agent 精修：输出适合直接阅读或交付的精华摘要。")
    return save_document(document, path, overwrite)


def create_corrections_docx(
    path: Path,
    transcript_path: Path,
    source_file: str,
    generated_at: str,
    output_formats: list[str],
    overwrite: bool,
) -> bool:
    Document, Pt, qn = import_docx()
    document = Document()
    set_doc_style(document, Pt, qn)
    document.add_heading("勘误与精修版", level=0)
    add_info_table(document, build_info_lines(transcript_path, source_file, generated_at, output_formats))
    document.add_heading("疑似识别错误表", level=1)
    add_simple_table(document, ["原句", "疑似错词", "建议修正", "理由/上下文", "置信度"], [["待 agent 核对", "待 agent 填写", "待 agent 填写", "结合上下文判断", "-"]])
    document.add_heading("专有名词统一表", level=1)
    add_simple_table(document, ["原句", "原转写", "建议修正", "类型", "置信度"], [["待 agent 核对", "待 agent 填写", "待 agent 填写", "人名/地名/机构名/品牌名/术语", "-"]])
    document.add_heading("语句精修说明", level=1)
    add_simple_table(document, ["原句", "建议修正", "类型", "理由/上下文", "置信度"], [["待 agent 核对", "待 agent 填写", "错词/断句/数字/术语/重复", "结合上下文判断", "-"]])
    document.add_heading("精修后正文", level=1)
    document.add_paragraph("待 agent 精修：在这里输出修正错别字、术语、断句和段落结构后的版本。")
    return save_document(document, path, overwrite)


def create_markdown_docx(path: Path, title: str, info_lines: list[str], markdown_text: str, overwrite: bool) -> bool:
    Document, Pt, qn = import_docx()
    document = Document()
    set_doc_style(document, Pt, qn)
    document.add_heading(title, level=0)
    add_info_table(document, info_lines)
    for block in parse_markdown_blocks(markdown_text):
        if block.kind == "heading":
            if block.level == 1 and block.text == title:
                continue
            document.add_heading(block.text, level=min(max(block.level, 1), 4))
        elif block.kind == "paragraph":
            for paragraph in [part.strip() for part in block.text.splitlines() if part.strip()]:
                document.add_paragraph(paragraph)
        elif block.kind == "list":
            for item in block.items:
                document.add_paragraph(item, style="List Bullet")
        elif block.kind == "table" and block.rows:
            add_simple_table(document, list(block.rows[0]), [list(row) for row in block.rows[1:]])
        elif block.kind == "code":
            document.add_paragraph(block.text)
    return save_document(document, path, overwrite)


def review_kind_from_markdown(path: Path) -> str | None:
    name = path.name
    if name.endswith(".summary.md"):
        return "summary"
    if name.endswith(".corrections.md"):
        return "corrections"
    if name.endswith(".publish.md"):
        return "publish"
    if name.endswith(".meeting-notes.md"):
        return "meeting-notes"
    return None


def base_stem_from_review_markdown(path: Path) -> str:
    name = path.name
    if name.endswith(".summary.md"):
        return name[: -len(".summary.md")]
    if name.endswith(".corrections.md"):
        return name[: -len(".corrections.md")]
    if name.endswith(".publish.md"):
        return name[: -len(".publish.md")]
    if name.endswith(".meeting-notes.md"):
        return name[: -len(".meeting-notes.md")]
    return path.stem


def title_for_kind(kind: str) -> str:
    titles = {
        "summary": "内容总结",
        "corrections": "勘误与精修版",
        "publish": "刊物发布版",
        "meeting-notes": "会议纪要版",
    }
    return titles.get(kind, "整理稿")


def sync_markdown_file(
    markdown_path: Path,
    output_dir: Path,
    include_docx: bool,
    include_html: bool,
    overwrite: bool,
) -> tuple[list[tuple[Path, bool]], list[str]]:
    if not markdown_path.exists():
        return [], [f"Markdown file not found: {markdown_path}"]
    kind = review_kind_from_markdown(markdown_path)
    if not kind:
        return [], [f"Expected a *.summary.md or *.corrections.md file: {markdown_path}"]

    markdown_text = read_text(markdown_path)
    base_stem = base_stem_from_review_markdown(markdown_path)
    title = markdown_title(markdown_text, title_for_kind(kind))
    generated_at = dt.datetime.now().astimezone().isoformat(timespec="seconds")
    output_formats = []
    if include_docx:
        output_formats.append(f"{kind}.docx")
    if include_html:
        output_formats.append(f"{kind}.html")
    info_lines = build_sync_info_lines(markdown_path, markdown_text, generated_at, output_formats)
    output_dir.mkdir(parents=True, exist_ok=True)
    results: list[tuple[Path, bool]] = []
    warnings: list[str] = []

    if include_docx:
        docx_path = output_dir / f"{base_stem}.{kind}.docx"
        try:
            results.append((docx_path, create_markdown_docx(docx_path, title, info_lines, markdown_text, overwrite)))
        except RuntimeError as exc:
            warnings.append(str(exc))
        except Exception as exc:  # noqa: BLE001 - syncing HTML should still succeed if DOCX fails.
            warnings.append(f"Could not sync Word output for {markdown_path}: {exc}")

    if include_html:
        html_path = output_dir / f"{base_stem}.{kind}.html"
        results.append(
            (
                html_path,
                write_if_needed(html_path, build_markdown_sync_html(title, info_lines, markdown_text), overwrite),
            )
        )

    return results, warnings


def sibling_review_markdown_paths(transcript_path: Path, output_dir: Path) -> list[Path]:
    stem = transcript_path.stem
    return [
        output_dir / f"{stem}.summary.md",
        output_dir / f"{stem}.corrections.md",
        output_dir / f"{stem}.publish.md",
        output_dir / f"{stem}.meeting-notes.md",
    ]


def sync_review_outputs(
    input_path: Path,
    output_dir: Path | None,
    include_docx: bool,
    include_html: bool,
    overwrite: bool,
) -> tuple[list[tuple[Path, bool]], list[str]]:
    resolved_output_dir = output_dir or input_path.parent
    results: list[tuple[Path, bool]] = []
    warnings: list[str] = []

    if review_kind_from_markdown(input_path):
        synced, sync_warnings = sync_markdown_file(input_path, resolved_output_dir, include_docx, include_html, overwrite)
        return synced, sync_warnings

    if not input_path.exists():
        return [], [f"Input file not found: {input_path}"]

    candidates = sibling_review_markdown_paths(input_path, resolved_output_dir)
    existing_candidates = [path for path in candidates if path.exists()]
    if not existing_candidates:
        return [], [f"No review Markdown files found beside: {input_path}"]

    for markdown_path in existing_candidates:
        synced, sync_warnings = sync_markdown_file(markdown_path, resolved_output_dir, include_docx, include_html, overwrite)
        results.extend(synced)
        warnings.extend(sync_warnings)

    return results, warnings


def selected_kinds(values: list[str] | None) -> list[str]:
    if values:
        return values
    return ["summary", "corrections"]


def build_markdown_for_kind(
    kind: str,
    transcript_path: Path,
    source_file: str,
    transcript: str,
    generated_at: str,
) -> str:
    if kind == "summary":
        return build_summary(transcript_path, source_file, transcript, generated_at)
    if kind == "corrections":
        return build_corrections(transcript_path, source_file, transcript, generated_at)
    if kind == "publish":
        return build_publish(transcript_path, source_file, transcript, generated_at)
    if kind == "meeting-notes":
        return build_meeting_notes(transcript_path, source_file, transcript, generated_at)
    raise ValueError(f"Unsupported deliverable kind: {kind}")


def markdown_filename(stem: str, kind: str) -> str:
    return f"{stem}.{kind}.md"


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Create transcript, summary, correction, publish, or meeting-note deliverables."
    )
    parser.add_argument("transcript", help="Path to a transcript file or review Markdown file.")
    parser.add_argument("--source-file", default="", help="Original media file path.")
    parser.add_argument("--output-dir", default=None, help="Directory for review outputs.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing review outputs.")
    parser.add_argument(
        "--transcript-only",
        action="store_true",
        help="Generate only the transcript Word file. This is the default transcribe flow.",
    )
    parser.add_argument(
        "--kind",
        action="append",
        choices=["summary", "corrections", "publish", "meeting-notes"],
        help="Review deliverable to create. Repeat for multiple kinds. Defaults to summary and corrections.",
    )
    parser.add_argument("--docx", action="store_true", help="Generate Word outputs. This is enabled by default.")
    parser.add_argument("--no-docx", action="store_true", help="Skip Word outputs.")
    parser.add_argument("--html", action="store_true", help="Generate HTML outputs for selected review deliverables.")
    parser.add_argument("--all", action="store_true", help="Generate Markdown, Word, and HTML outputs.")
    parser.add_argument("--markdown-only", action="store_true", help="Generate only Markdown outputs.")
    parser.add_argument(
        "--sync",
        action="store_true",
        help="Read existing review Markdown and sync its current content to DOCX/HTML.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    transcript_path = expand_path(args.transcript)
    output_dir = expand_path(args.output_dir) if args.output_dir else default_output_dir()
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        include_docx = not args.no_docx and not args.markdown_only
        include_html = (args.html or args.all) and not args.markdown_only

        if args.sync:
            if not args.html and not args.all and not args.markdown_only:
                include_html = True
            results, warnings = sync_review_outputs(
                input_path=transcript_path,
                output_dir=output_dir if args.output_dir else None,
                include_docx=include_docx,
                include_html=include_html,
                overwrite=True if not args.overwrite else args.overwrite,
            )
            print("Review sync outputs:")
            for path, written in results:
                print(f"  {output_status(path, written)}: {path}")
            for warning in warnings:
                print(f"Warning: {warning}", file=sys.stderr)
            if not results and warnings:
                return 1
            print_sync_guidance(results)
            return 0

        if not transcript_path.exists():
            raise FileNotFoundError(f"Transcript not found: {transcript_path}")
        segments = read_transcript_segments(transcript_path)
        transcript = "\n".join(segment.text for segment in segments if segment.text).strip()
        stem = transcript_path.stem
        generated_at = dt.datetime.now().astimezone().isoformat(timespec="seconds")
        kinds = selected_kinds(args.kind)
        output_formats = []
        if args.transcript_only:
            output_formats.append("transcript.docx")
        else:
            output_formats.extend(f"{kind}.md" for kind in kinds)
        if include_docx:
            if "transcript.docx" not in output_formats:
                output_formats.append("transcript.docx")
            if not args.transcript_only:
                output_formats.extend(f"{kind}.docx" for kind in kinds)
        if include_html:
            output_formats.extend(f"{kind}.html" for kind in kinds)

        transcript_docx_path = output_dir / f"{stem}.transcript.docx"

        results: list[tuple[Path, bool]] = []
        warnings: list[str] = []

        if include_docx:
            try:
                results.append(
                    (
                        transcript_docx_path,
                        create_transcript_docx(
                            transcript_docx_path,
                            transcript_path,
                            args.source_file,
                            generated_at,
                            output_formats,
                            segments,
                            args.overwrite,
                        ),
                    )
                )
            except RuntimeError as exc:
                warnings.append(str(exc))
            except Exception as exc:  # noqa: BLE001 - review deliverables should not break transcription.
                warnings.append(f"Could not create transcript Word output: {exc}")

        if not args.transcript_only:
            for kind in kinds:
                markdown_path = output_dir / markdown_filename(stem, kind)
                markdown_text = build_markdown_for_kind(kind, transcript_path, args.source_file, transcript, generated_at)
                results.append((markdown_path, write_if_needed(markdown_path, markdown_text, args.overwrite)))

                if include_docx:
                    docx_path = output_dir / f"{stem}.{kind}.docx"
                    try:
                        if kind == "summary":
                            written = create_summary_docx(
                                docx_path,
                                transcript_path,
                                args.source_file,
                                generated_at,
                                output_formats,
                                transcript,
                                args.overwrite,
                            )
                        elif kind == "corrections":
                            written = create_corrections_docx(
                                docx_path,
                                transcript_path,
                                args.source_file,
                                generated_at,
                                output_formats,
                                args.overwrite,
                            )
                        else:
                            info_lines = build_info_lines(transcript_path, args.source_file, generated_at, output_formats)
                            written = create_markdown_docx(
                                docx_path,
                                title_for_kind(kind),
                                info_lines,
                                markdown_text,
                                args.overwrite,
                            )
                        results.append((docx_path, written))
                    except RuntimeError as exc:
                        warnings.append(str(exc))
                    except Exception as exc:  # noqa: BLE001 - other deliverables should still be created.
                        warnings.append(f"Could not create Word output for {kind}: {exc}")

                if include_html:
                    html_path = output_dir / f"{stem}.{kind}.html"
                    info_lines = build_info_lines(transcript_path, args.source_file, generated_at, output_formats)
                    if kind == "summary":
                        html_text = build_summary_html(info_lines, transcript)
                    elif kind == "corrections":
                        html_text = build_corrections_html(info_lines)
                    else:
                        html_text = build_markdown_sync_html(title_for_kind(kind), info_lines, markdown_text)
                    results.append((html_path, write_if_needed(html_path, html_text, args.overwrite)))

        print("Transcript outputs:" if args.transcript_only else "Review draft outputs:")
        for path, written in results:
            print(f"  {output_status(path, written)}: {path}")
        if warnings:
            for warning in warnings:
                print(f"Warning: {warning}", file=sys.stderr)
        if args.transcript_only:
            print_transcript_guidance(transcript_path, output_dir, stem)
        else:
            print_initial_review_guidance(transcript_path, output_dir, stem, kinds, include_html)
        return 0
    except Exception as exc:  # noqa: BLE001 - CLI should print friendly errors.
        print(f"Error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
